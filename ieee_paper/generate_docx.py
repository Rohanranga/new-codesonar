"""
CodeSonar IEEE Paper — proper two-section Word document.
Section 1 (single-column): title + authors + separator
Section 2 (two-column):    abstract + body + references
Run: python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ── Twip constants (1 cm = 567 twips) ────────────────────────────────────────
def cm2tw(cm): return str(int(cm * 567))

A4_W    = cm2tw(21.0)   # page width
A4_H    = cm2tw(29.7)   # page height
MAR_L   = cm2tw(1.8)
MAR_R   = cm2tw(1.8)
MAR_T   = cm2tw(2.0)
MAR_B   = cm2tw(2.2)
COL_GAP = "720"         # 0.5 inch gap between two columns

# Author column centers (from left margin, full 17.4 cm usable width)
COL1 = int(2.9  * 567)
COL2 = int(8.7  * 567)
COL3 = int(14.5 * 567)

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()
sec0 = doc.sections[0]
sec0.page_width    = Cm(21.0)
sec0.page_height   = Cm(29.7)
sec0.left_margin   = Cm(1.8)
sec0.right_margin  = Cm(1.8)
sec0.top_margin    = Cm(2.0)
sec0.bottom_margin = Cm(2.2)

# ── Font helper ───────────────────────────────────────────────────────────────
def sf(run, size=10, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic

# ── Paragraph helpers ─────────────────────────────────────────────────────────
def add_para(text="", align=WD_ALIGN_PARAGRAPH.LEFT,
             size=10, bold=False, italic=False, sb=0, sa=0, ls=13):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(ls)
    if text:
        r = p.add_run(text)
        sf(r, size=size, bold=bold, italic=italic)
    return p

def body(text, sa=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(11.5)
    r = p.add_run(text)
    sf(r, size=10)
    return p

def sec_head(roman, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    label = f"{roman}.   {title.upper()}" if roman else title.upper()
    r = p.add_run(label)
    sf(r, size=10, bold=True)

def sub_head(letter, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(f"{letter}. {title}")
    sf(r, size=10, bold=True)

def eq(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    sf(r, size=10.5, italic=True)

def fig(img_path, num, caption, w=7.2):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        p.add_run().add_picture(img_path, width=Cm(w))
    except Exception:
        body(f"[Figure {num}: {caption}]")
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(1)
    cp.paragraph_format.space_after  = Pt(4)
    r = cp.add_run(f"Fig. {num}. {caption}")
    sf(r, size=8.5, italic=True)

def numbered_item(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(11.5)
    r = p.add_run(text)
    sf(r, size=10)

# Removed tab-stop helpers

# ── Obsolete section helpers removed ──────────────────────────────────────────

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Single-column: Title + Authors
# ═══════════════════════════════════════════════════════════════════════════════

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(0)
tp.paragraph_format.space_after  = Pt(18)
r = tp.add_run(
    "Code Comprehension and Debugging Platform\n"
    "Using RAG"
)
sf(r, size=20, bold=False)  # Not bold in reference image, just larger text

authors_data = [
    ("Karnala Akaniksha",
     "Computer Science and Engineering (AI\n& ML)",
     "Nalla Malla Reddy Engineering\nCollege",
     "Hyderabad, India",
     "karnalaakaniksha11@nmrec.edu.in"),
    ("Kommu Tharun Kumar",
     "Computer Science and Engineering (AI\n& ML)",
     "Nalla Malla Reddy Engineering\nCollege",
     "Hyderabad, India",
     "tharun28k@gmail.com"),
    ("Ajmeera Rakesh Kumar",
     "Computer Science and Engineering (AI\n& ML)",
     "Nalla Malla Reddy Engineering\nCollege",
     "Hyderabad, India",
     "ajmeerarakeshkumar@gmail.com"),
    ("Ranga Rohan",
     "Computer Science and Engineering (AI\n& ML)",
     "Nalla Malla Reddy Engineering\nCollege",
     "Hyderabad, India",
     "uniquerohankumar@gmail.com"),
    ("Baswapathiruni Krishna Kumar",
     "Computer Science and Engineering (AI\n& ML)",
     "Nalla Malla Reddy Engineering\nCollege",
     "Hyderabad, India",
     "krishnakumar.cse@nmrec.edu.in"),
    ("S. Ramchandra Reddy",
     "Computer Science and Engineering (AI\n& ML)",
     "Nalla Malla Reddy Engineering\nCollege",
     "Hyderabad, India",
     "rcreddy79@gmail.com"),
]

auth_table = doc.add_table(rows=2, cols=3)
auth_table.autofit = False

# Remove borders
for row in auth_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            border_el = OxmlElement(f'w:{side}')
            border_el.set(qn('w:val'), 'none')
            tcBorders.append(border_el)
        tcPr.append(tcBorders)

for idx, auth in enumerate(authors_data):
    row_i = idx // 3
    col_i = idx % 3
    cell = auth_table.rows[row_i].cells[col_i]
    cell.width = Cm(5.8)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8) if row_i > 0 else Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    
    for i, line in enumerate(auth):
        run = p.add_run(line + ("\n" if i < len(auth)-1 else ""))
        sf(run, size=9, italic=True if i < 4 else False)

# Empty paragraph acting as the boundary before the section break
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(8)
sep.paragraph_format.space_after  = Pt(0)

# Create Section 2 (CONTINUOUS style) for two-column flow
new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
sectPr = new_sec._sectPr
cols = OxmlElement('w:cols')
cols.set(qn('w:num'), '2')
cols.set(qn('w:space'), '720')
cols.set(qn('w:equalWidth'), '1')
sectPr.append(cols)

new_sec.page_width    = Cm(21.0)
new_sec.page_height   = Cm(29.7)
new_sec.left_margin   = Cm(1.8)
new_sec.right_margin  = Cm(1.8)
new_sec.top_margin    = Cm(2.0)
new_sec.bottom_margin = Cm(2.2)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Two-column: Abstract + full paper body
# ═══════════════════════════════════════════════════════════════════════════════

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_abs.paragraph_format.space_before = Pt(1)
p_abs.paragraph_format.space_after  = Pt(3)
p_abs.paragraph_format.line_spacing = Pt(11.5)
r1 = p_abs.add_run("Abstract")
sf(r1, size=9, bold=True, italic=True)
r2 = p_abs.add_run(
    "\u2014Getting to grips with a codebase you've never seen before is one of those "
    "tasks that sounds simpler than it actually is. You pull the repository, get it "
    "running, and then you're left completely on your own to figure out what any "
    "of it actually does. This paper describes CodeSonar, a web-based tool we built "
    "specifically to shorten that disorienting gap. A developer pastes a GitHub URL, "
    "or uploads a zip file, and within a minute or so gets back a plain-English "
    "summary of the whole project, an explanation of each file, a diagram of how the "
    "modules relate to each other, a rundown of possible bugs and security gaps, and "
    "a chat panel for whatever follow-up questions they still have. Under the hood, "
    "it uses Retrieval-Augmented Generation\u2014ranking code segments by relevance "
    "before handing any of it to Google\u2019s Gemini 2.5 Flash model\u2014so the "
    "responses stay grounded in the actual repository. Nothing is stored on our end, "
    "which sidesteps the usual worry about sending proprietary code to an external "
    "service. We ran the tool on ten open-source projects of different sizes and found "
    "that the explanations rated noticeably better than generic LLM output, especially "
    "on projects with sparse or absent documentation."
)
sf(r2, size=9)

p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_kw.paragraph_format.space_after  = Pt(3)
p_kw.paragraph_format.line_spacing = Pt(11.5)
rk1 = p_kw.add_run("Keywords\u2014")
sf(rk1, size=9, bold=True, italic=True)
rk2 = p_kw.add_run(
    "code comprehension, GitHub repository analysis, Retrieval-Augmented "
    "Generation, Gemini 2.5 Flash, multi-agent LLM, architecture visualization, "
    "complexity analysis, AI-assisted debugging"
)
sf(rk2, size=9, italic=True)

# ── I. INTRODUCTION ───────────────────────────────────────────────────────────
sec_head("I", "Introduction")

body(
    "Anyone exploring open-source software quickly runs into a familiar roadblock: you clone a project, "
    "get it to compile, and then stare blankly at the folder structure trying to guess how "
    "everything wires together. README files almost entirely focus on setup steps. Code comments, "
    "if they exist, usually speak to people who are already deeply involved in the project. "
    "Unit tests reveal what functions should return, but rarely explain why the architecture was "
    "designed that way in the first place. This leaves newcomers feeling like they showed up "
    "to a movie during the third act, struggling to piece the plot together without any outside help."
)
body(
    "A common misconception is that this confusion mainly affects junior programmers. In reality, "
    "even seasoned engineers can lose days trying to build a mental map of a sprawling backend service "
    "they just inherited. The core difficulty originates from the sheer volume of material. "
    "With platforms like GitHub hosting massive numbers of repositories, nobody has enough time "
    "to manually trace every function call just to evaluate if a codebase is worth using or modifying."
)
body(
    "Of course, we have traditional static analysis. Linters complain about messy syntax, "
    "and scanners loudly warn us about outdated dependencies. Yet all of this happens at a purely "
    "mechanical level. A linter cannot look at a messy authentication controller and say, 'Hey, "
    "this handles user logins but it's dangerously coupled to the database config.' To get that "
    "kind of insight, you traditionally had to bother the original developer. "
    "Recently, large language models (LLMs) proved they can do some of this explaining for us. "
    "However, if you blindly dump an entire modern repository into a chatbot, the system quickly chokes on the token limits."
)
body(
    "We built CodeSonar to navigate around exactly that limitation. Rather than shoving thousands "
    "of files into an LLM and crossing our fingers, the system aggressively scores and filters "
    "code chunks down to just the highly relevant pieces before asking the AI any questions. "
    "By adopting a Retrieval-Augmented Generation (RAG) methodology for repository exploration, "
    "we managed to create a web platform that takes a fresh GitHub link and returns a full, "
    "human-readable technical breakdown in under two minutes."
)

# ── II. RELATED WORK ──────────────────────────────────────────────────────────
sec_head("II", "Related Work")

body(
    "The tools developers already rely on for code quality\u2014SonarQube, ESLint, "
    "PMD\u2014are good at what they do, but they stay at the surface level. They "
    "check syntax, flag known anti-patterns, and count complexity metrics. What they "
    "cannot do is read a function, understand what business requirement it satisfies, "
    "and explain that in plain language. Documentation generators like Doxygen and "
    "TypeDoc get closer but rely entirely on the quality of comments already written "
    "into the source. When those comments are missing, which they often are, the "
    "generated output is little more than a formatted list of method signatures."
)
body(
    "The arrival of large language models opened a new possibility. Nam et al. [11] "
    "showed that LLMs can meaningfully help with code understanding when used "
    "interactively, though accuracy dropped noticeably on questions that required "
    "context from multiple files at once. Mansur et al. [4] addressed a related "
    "problem\u2014automated code repair\u2014by retrieving Stack Overflow posts as "
    "context before asking the model to suggest a fix. Their results confirmed "
    "something that is now fairly well understood: an LLM with good context produces "
    "far better answers than an LLM without it."
)
body(
    "What has been explored less is whether the same principle holds for whole-"
    "repository comprehension rather than isolated bug fixes. Stojanovic et al. [10] "
    "built a multi-agent system for software documentation and found that splitting "
    "the task across specialized agents produced qualitatively better output than a "
    "single large prompt. Ramasamy et al. [19] analysed dependency structures across "
    "public repositories and showed that structural patterns can be inferred "
    "automatically from import graphs\u2014a technique we use directly in our "
    "architecture diagram generator. Swacha and Gracel [17] surveyed RAG chatbot "
    "applications in education and found high satisfaction when responses were grounded "
    "in specific material rather than general model knowledge."
)
body(
    "Taken together, this body of work points toward the same conclusion we arrived "
    "at independently: a system that retrieves relevant code and grounds every answer "
    "in actual repository content will outperform both generic LLM prompting and "
    "traditional static analysis. What had not been built was an integrated, "
    "browser-accessible platform combining retrieval, multi-agent inference, "
    "visualization, an interactive chatbot, and privacy-preserving real-time "
    "processing in one place. That is what CodeSonar provides."
)

# ── III. SYSTEM ARCHITECTURE ──────────────────────────────────────────────────
sec_head("III", "System Architecture")

body(
    "We chose Next.js 14 partly out of familiarity and partly because its App Router "
    "lets you collocate API handlers with the frontend without a separate backend "
    "service. The overall structure, shown in Fig. 1, has three layers. The browser "
    "side is a React application that manages input and renders results. Two server "
    "API routes handle the heavy work: /api/analyze runs the full pipeline when a "
    "repository is submitted, and /api/chat handles conversational follow-up. All "
    "Gemini API calls go through the server so the key never reaches the browser."
)
body(
    "We settled on Gemini 2.5 Flash after trying a few alternatives. It has a large "
    "enough context window for our needs, responds fast enough that users don't give "
    "up waiting, and is priced reasonably for a project on a student budget. "
    "Temperature is set to 0.4 across all five agents\u2014low enough to keep factual "
    "claims accurate, high enough that the summaries read like prose rather than "
    "bullet-pointed output."
)
body(
    "One decision we made early and stuck with: no persistent storage. Repository "
    "files are fetched, processed in memory, and then the results are returned. "
    "Once the response is sent, nothing is retained on the server. This matters "
    "because most of the codebases developers want to understand include code they "
    "wouldn't want sitting on someone else's servers. The tradeoff is that running "
    "the same repository twice means running the pipeline twice, but with the whole "
    "thing finishing under ninety seconds that hasn't been a real complaint."
)

fig(r"R:\dup\codesonar\ieee_paper\figures\architecture.png",
    1, "Three-Tier Architecture of the CodeSonar Platform", 7.5)

# ── IV. METHODOLOGY ───────────────────────────────────────────────────────────
sec_head("IV", "Methodology")

body(
    "The pipeline that runs when a URL is submitted goes through five distinct stages. "
    "None of them is especially complicated on its own; the value is in running them "
    "in order and having each stage feed cleanly into the next."
)

sub_head("A", "Fetching the Repository")
body(
    "Submitting a URL triggers a recursive walk of the GitHub Contents API. Every "
    "file in the tree gets fetched and held in memory as a path-to-content map. "
    "Binary files and anything over a megabyte are skipped\u2014those are usually assets, "
    "not code\u2014and the same logic applies when someone uploads a file directly "
    "rather than providing a URL."
)

sub_head("B", "Cleaning Up Before Analysis")
body(
    "A typical JavaScript repository comes with its node_modules directory, several "
    "lockfiles, compiled output, source maps, and minified bundles. None of those tell "
    "you anything useful about how the project works. Our filtering pass removes "
    "anything whose path includes node_modules or lock, or whose extension ends "
    "in .min.js or .map. What remains is what a developer would actually sit and "
    "read. The total size of that cleaned set is:"
)
eq("LOC_total  =  \u03a3 (i = 1 to n)  LOC_i")
body(
    "where LOC_i is simply how many lines the i-th file has, and n stands for the total "
    "number of files we decided to keep. This count gets prominently displayed to the user "
    "and acts as a baseline metric for our later complexity heuristics."
)

sub_head("C", "Per-File Complexity Estimate")
body(
    "Before handing anything over to the LLM agents, we run a rapid local heuristic on every single file. "
    "We essentially just tally up the function declarations F and class definitions K, merging them like so:"
)
eq("C_f  =  \u03b1F  +  \u03b2K")
body(
    "For our purposes, we explicitly set \u03b1\u202f=\u202f1.0 and \u03b2\u202f=\u202f2.0. We penalize classes "
    "more heavily simply because they tend to carry messy state and inheritance chains, which naturally forces "
    "readers to juggle more details in their head. Whenever this equation pushes a file's score too high, "
    "our dashboard flags it as a prime target for future refactoring efforts."
)

sub_head("D", "RAG Context Selection")
body(
    "The hardest part of this entire pipeline is isolating exactly which snippets of code to actually feed into the prompt. "
    "We tackle this by calculating the cosine similarity across TF-IDF vectors:"
)
eq("R(q, d)  =  (q \u00b7 d)  /  ( \u2016q\u2016 \u00b7 \u2016d\u2016 )")
body(
    "where q acts as the weighted term vector for whatever the AI agent is asking, while d represents the vector "
    "of a specific code chunk. By filtering down to just the top eight matches, we ensure the LLM gets enough "
    "immediate context to answer intelligently, without overwhelming the model's attention span with irrelevant noise."
)

sub_head("E", "The Agent Chain")
body(
    "Running everything as one large prompt gave broad but shallow output. Splitting "
    "the work into five specialized agents, each with a single clearly stated task, "
    "gave noticeably better results."
)
body(
    "The Tech Stack Agent inspects imports, config files, and extensions to produce "
    "a list of every language, framework, and build tool in the project. It's fairly "
    "reliable at separating development dependencies from runtime ones."
)
body(
    "The Complexity Agent looks for high coupling between modules, obvious duplication, "
    "oversized functions, and test-coverage signals. It outputs a 1\u201310 score plus "
    "a plain-English justification\u2014the kind of read-out a senior engineer might "
    "give after an initial pass through the code."
)
body(
    "The Improvement Agent works like a code reviewer. It finds specific issues tied "
    "to real file paths and line ranges rather than offering vague advice, and "
    "categorizes each finding as a security problem, a performance issue, or a "
    "refactoring opportunity."
)
body(
    "The Bug and Explanation Agent does two things in one pass because both require "
    "careful line-level reading. It scans for security vulnerabilities\u2014"
    "unsanitized inputs, hardcoded secrets, missing access controls\u2014while also "
    "writing a plain-English description of every file's purpose. Combining them "
    "saved one API call per file, which matters when you're working within a "
    "rate limit of fifteen requests per minute."
)
body(
    "The Summary Agent runs last. It reads everything the other four produced and "
    "writes a few paragraphs describing the project as a whole: what it is, how "
    "it's structured, what the code quality is like, and where a new contributor "
    "should start. That summary appears at the top of the dashboard and is what "
    "most people read first."
)
body("Every agent call follows the same pattern:")
eq("Output  =  LLM( Query + Context )")
body(
    "All five agents share a single Gemini instance and run sequentially to stay "
    "within the rate limit. A shared retry helper catches 429 responses, backs off "
    "exponentially starting at eight seconds, and gives up after three attempts."
)

sub_head("F", "Architecture Diagram")
body(
    "While the AI agents are busy running their analysis, a completely parallel backend process scans through the "
    "project's import syntax (like 'require' or 'import' statements). It constructs a directed network graph G\u202f=\u202f(V,\u202fE), "
    "meaning every single file acts as a vertex while their static dependencies form the edges. "
    "We then convert this mathematical graph directly into a Mermaid.js diagram so the frontend can draw an interactive map. "
    "Honestly, during informal testing, users loved this feature more than almost anything else. Physically seeing the arrows "
    "drawn between modules makes a project so much easier to grasp compared to just reading text."
)

sub_head("G", "End-to-End Pipeline Steps")
for step in [
    "A user pastes a URL or drags a zip file into the landing page.",
    "Our server reaches out via the GitHub API and downloads the raw tree.",
    "A cleanup script aggressively strips out junk like lockfiles and minified bundles.",
    "The surviving code is chopped into chunks and mathematically ranked for relevance.",
    "The Tech Stack AI runs first, swiftly followed by the Complexity AI.",
    "Next up, the Improvement AI and Bug Explanation agents process the code.",
    "While that happens, the system translates the file imports into a Mermaid graph.",
    "A final Summary AI reads everything produced so far and writes an executive overview.",
    "The React frontend receives the final JSON bundle and renders the dashboard.",
]:
    numbered_item(step)

fig(r"R:\dup\codesonar\ieee_paper\figures\workflow.png",
    2, "End-to-End Workflow of the CodeSonar Analysis Pipeline", 7.5)
fig(r"R:\dup\codesonar\ieee_paper\figures\dataflow.png",
    3, "RAG-Based Data Flow from User Query to AI Output", 7.5)

# ── V. RESULTS AND DISCUSSION ─────────────────────────────────────────────────
sec_head("V", "Results and Discussion")

body(
    "To figure out if CodeSonar actually worked in the real world, we threw ten totally different public repositories at it. "
    "We purposely picked a chaotic mix: a few Next.js frontend templates, some Node microservices, a couple of heavy Spring Boot "
    "business apps, and even some Python tooling. They ranged anywhere from a couple thousand lines to a surprisingly dense forty-five "
    "thousand line behemoth. We then grabbed three developers who had literally never seen these projects before, and asked them to "
    "grade the AI's explanations on a basic 1 to 5 scale for accuracy, depth, and overall completeness."
)

# TABLE I
tc = doc.add_paragraph()
tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
tc.paragraph_format.space_before = Pt(6)
tc.paragraph_format.space_after  = Pt(2)
r_tc = tc.add_run("TABLE I.   Evaluation Results Across Repository Sizes")
sf(r_tc, size=9, bold=True, italic=True)

tbl = doc.add_table(rows=4, cols=5)
tbl.style = 'Table Grid'
for ci, h in enumerate(["Repository Size", "Accuracy\n(avg/5)",
                          "Specificity\n(avg/5)", "Completeness\n(avg/5)", "Avg. Time"]):
    cell = tbl.rows[0].cells[ci]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(h)
    sf(r, size=9, bold=True)
for ri, row in enumerate([
    ["Small (< 5 K LOC)",   "4.6", "4.5", "4.4", "18 s"],
    ["Medium (5–20 K LOC)", "4.3", "4.1", "4.2", "42 s"],
    ["Large (> 20 K LOC)",  "3.9", "3.8", "3.7", "89 s"],
]):
    for ci, val in enumerate(row):
        cell = tbl.rows[ri + 1].cells[ci]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(val)
        sf(r, size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

body(
    "Glancing at Table I, the data essentially confirms what we initially suspected. Tiny projects got excellent grades "
    "because the model could comfortably cram a huge percentage of the repository into its immediate memory. "
    "The larger projects obviously scored lower, but frankly, they held up way better than our team anticipated. "
    "Even when dealing with the massive 45k-line repository, the scores hovered surprisingly close to a 4.0. "
    "This clearly indicates that our similarity ranking logic is doing a solid job of fishing out the right files "
    "even when 90% of the codebase has to be ignored."
)
body(
    "One of the biggest surprises was the vulnerability scanner. Between our hardcoded heuristic rules and the "
    "LLM's ability to read between the lines, the system spotted eighty-seven distinct issues. What's crazy is that "
    "over sixty of those were totally undocumented bugs that nobody had reported on GitHub yet. "
    "We saw tons of missing input sanitization and developers hardcoding API keys right into their config files. "
    "When we ran normal, rigid linters against the exact same folders, they barely caught thirty-four problems, "
    "which really highlights why having an AI that actually 'reads' intent is so powerful."
)
body(
    "As for the visual architecture diagrams, the feedback was overwhelmingly positive. "
    "Almost every single tester mentioned that they kept switching back to the visual flowchart while reading the "
    "textual breakdowns just to keep themselves oriented. The only genuine complaint we got was regarding that massive "
    "Python project; rendering a spiderweb of sixty-plus connected files into a single browser window was undeniably "
    "overwhelming and chaotic. We definitely need to figure out a better way to collapse or group those visuals in the future."
)
body(
    "The interactive chat box also proved way more popular than we thought it would be. "
    "People wouldn't just read the summary and leave; they would start interrogating the bot with highly specific queries like "
    "'Wait, which exact folder handles the JWT tokens?' Because our chatbot uses the underlying RAG index instead of just "
    "pulling randomly from the internet, it rarely hallucinated. Rather than giving vague boilerplate advice, it would spit "
    "out the exact file names."
)

# ── VI. CONCLUSION AND FUTURE WORK ───────────────────────────────────────────
sec_head("VI", "Conclusion and Future Work")

body(
    "The whole reason our team started this project is because we were frankly tired of wasting our weekends "
    "staring at undocumented code trying to reverse-engineer how a web app worked. We knew LLMs were smart, "
    "but they needed highly curated context to stop hallucinating. Bringing together RAG similarity searches "
    "and a pipeline of specialized AI agents turned out to be the exact right formula. Now, developers can paste "
    "a random repository link and get a beautifully formatted, genuinely accurate breakdown of the software's architecture "
    "without ever having to spin up a local server."
)
body(
    "Looking at our initial testing metrics, we are quite optimistic. Holding a 3.7+ accuracy score across heavily "
    "complex repositories proves that context-injection beats naive prompting every time. Moreover, catching bugs "
    "that traditional linters completely miss showcases the immense potential of semantic code analysis. "
    "Of course, the system definitely has rough edges. Wait times creeping over a minute for huge repositories can test "
    "a programmer's patience, and those gigantic Mermaid diagrams desperately need a way to dynamically collapse nodes."
)
body(
    "Moving forward, our immediate priority has to be supporting private repositories. Asking developers to only scan "
    "public code is wildly restrictive, so wiring up proper OAuth integration for enterprise GitLab and Bitbucket accounts "
    "is next on our roadmap. We also want to empower users by adding a feature where the AI doesn't just point out a bug, "
    "but actively writes a pull request to patch it. And finally, implementing a system where users can thumbs-down bad "
    "explanations would give us the exact raw data we need to continuously tune our models."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
sec_head("", "References")

refs = [
    "F. Ros\u00e9n and M. R\u00fcbner, \u201cAugmented large language models for software engineering,\u201d 2024.",
    "S. Bag et al., \u201cRAG beyond text: enhancing image retrieval in RAG systems,\u201d in Proc. 2024 Int. Conf. Electrical, Computer and Energy Technologies (ICECET), IEEE, 2024.",
    "P. H. Russell et al., \u201cA large-scale analysis of bioinformatics code on GitHub,\u201d PloS ONE, vol. 13, no. 10, p. e0205898, 2018.",
    "E. Mansur et al., \u201cRAGFix: Enhancing LLM code repair using RAG and Stack Overflow posts,\u201d in Proc. 2024 IEEE Int. Conf. Big Data, IEEE, 2024.",
    "A. D. SP et al., \u201cAnalysing code-based RAG methods for knowledge retention,\u201d in Proc. 2025 4th Int. Conf. Advances in Computing (ACCESS), IEEE, 2025.",
    "Q. Romero Lauro et al., \u201cRAG without the lag: interactive debugging for RAG pipelines,\u201d arXiv e-prints, 2025.",
    "M. Antal and K. Buza, \u201cEvaluating open-source LLMs in RAG systems,\u201d Acta Universitatis Sapientiae, Informatica, vol. 17, no. 1, 2025.",
    "Y. Sugiyama et al., \u201cRelationship between model-based decision-making and source code comprehension,\u201d IEEE Trans. Software Eng., 2025.",
    "K. Ye, L. Su, and C. Qian, \u201cImportSnare: Directed code manual hijacking in RAG code generation,\u201d arXiv:2509.07941, 2025.",
    "D. Stojanovi\u0107 et al., \u201cUnit test generation multi-agent AI system for software documentation,\u201d in Proc. 2024 32nd Telecommunications Forum (TELFOR), IEEE, 2024.",
    "D. Nam et al., \u201cUsing an LLM to help with code understanding,\u201d in Proc. IEEE/ACM 46th Int. Conf. Software Engineering, 2024.",
    "P. Omrani et al., \u201cHybrid RAG approach for LLMs query response enhancement,\u201d in Proc. 2024 10th Int. Conf. Web Research (ICWR), pp. 22\u201326, IEEE, 2024.",
    "W. Yang et al., \u201cEnhancing the code debugging ability of LLMs via communicative agent based data refinement,\u201d language, vol. 30, p. 31, 2024.",
    "B. Sepidband et al., \u201cEnhancing LLM-based code generation with complexity metrics,\u201d arXiv:2505.23953, 2025.",
    "B. Tural, Z. \u00d6rpek, and Z. Destan, \u201cRAG and LLM integration,\u201d in Proc. 2024 8th Int. Symp. Innovative Approaches in Smart Technologies (ISAS), pp. 1\u20135, IEEE, 2024.",
    "R. Tian et al., \u201cDebugbench: Evaluating debugging capability of large language models,\u201d arXiv:2401.04621, 2024.",
    "J. Swacha and M. Gracel, \u201cRAG chatbots for education: a survey of applications,\u201d Applied Sciences, vol. 15, no. 8, p. 4234, 2025.",
    "J. Swacha and M. Gracel, \u201cRetrieval-augmented generation chatbots for education,\u201d Applied Sciences, vol. 15, no. 8, p. 4234, 2025.",
    "D. Ramasamy et al., \u201cWorkflow analysis of data science code in public GitHub repositories,\u201d Empirical Software Engineering, vol. 28, no. 1, p. 7, 2023.",
    "A. Lekssays et al., \u201cLLMxCPG: context-aware vulnerability detection through code property graph-guided LLMs,\u201d in Proc. 34th USENIX Security Symp., pp. 489\u2013507, 2025.",
]
for ref in refs:
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    r = p.add_run(ref)
    sf(r, size=8.5)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = r"R:\dup\codesonar\ieee_paper\CodeSonar_IEEE_Paper.docx"
doc.save(out)
print(f"Saved: {out}")
